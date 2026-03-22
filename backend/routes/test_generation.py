import io
import logging
from fastapi import APIRouter, Depends, File, HTTPException, UploadFile
from sqlalchemy.orm import Session
from config.database import get_db
from database.schemas import (
    GenerateTestRequest,
    GenerateTestLinkRequest,
    GenerateTestLinkResponse,
    BulkSendTestRequest,
)
from database.schemas import CandidateCreate
from database.models import JobDescription
from services.candidate_service import get_or_create_candidate
from services.test_service import create_test_session
from services.email_service import send_test_link_email

logger = logging.getLogger(__name__)
router = APIRouter()


def _clean_job_description(text: str) -> str:
    """
    Sanitize job description text before processing.

    If the text looks like raw binary .docx content (starts with 'PK' magic bytes or
    contains NUL bytes), attempt to extract readable text via python-docx.
    Falls back to stripping non-printable characters.
    Raises ValueError with HTTP 400 if no readable content can be extracted.
    """
    if "\x00" not in text and not text.startswith("PK"):
        return text

    # Attempt python-docx parsing
    try:
        import docx

        raw_bytes = text.encode("latin-1", errors="replace")
        doc = docx.Document(io.BytesIO(raw_bytes))
        extracted = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        if extracted.strip():
            return extracted
    except Exception:
        pass

    # Fallback: strip NUL and non-printable characters
    cleaned = "".join(ch for ch in text if ch >= " " or ch in "\n\r\t")
    if not cleaned.strip():
        raise ValueError(
            "Could not extract readable text from the uploaded file. "
            "Please paste the job description as plain text or upload a .txt file."
        )
    return cleaned



@router.post("/generate-test")
async def generate_test(request: GenerateTestRequest, db: Session = Depends(get_db)):
    """
    Generate interview questions from a job description.
    Delegates to the agents/generate_langgraph workflow and persists the JD.
    """
    try:
        from agents.generate_langgraph import run_generate_workflow

        cleaned_jd = _clean_job_description(request.job_description)
        result = run_generate_workflow(cleaned_jd)

        # Persist the job description (store cleaned text, not raw binary)
        jd = JobDescription(
            raw_text=cleaned_jd,
            role=result.get("role"),
            skills=result.get("skills"),
            weights=result.get("weights"),
            questions=result.get("questions"),
        )
        db.add(jd)
        db.commit()
        db.refresh(jd)

        result["jd_id"] = jd.id
        return result
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc))
    except ImportError:
        logger.warning("generate_langgraph not available, returning placeholder.")
        raise HTTPException(
            status_code=503,
            detail="Question generation workflow not configured. Please set up agents.",
        )
    except Exception as exc:
        logger.error("Test generation failed: %s", exc)
        raise HTTPException(status_code=500, detail=str(exc))


@router.post("/parse-file")
async def parse_file(file: UploadFile = File(...)):
    """
    Parse an uploaded .docx, .pdf, or .txt file and return its text content.
    The frontend calls this endpoint instead of reading binary files locally.
    """
    filename = (file.filename or "").lower()
    content = await file.read()

    try:
        if filename.endswith(".docx"):
            try:
                import docx

                doc = docx.Document(io.BytesIO(content))
                text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            except Exception as exc:
                raise HTTPException(
                    status_code=422,
                    detail=f"Could not parse .docx file: {exc}",
                )
        elif filename.endswith(".pdf"):
            try:
                import pdfplumber

                with pdfplumber.open(io.BytesIO(content)) as pdf:
                    text = "\n".join(
                        page.extract_text() or "" for page in pdf.pages
                    )
            except ImportError:
                # pdfplumber not installed; fall back to raw decode
                text = content.decode("utf-8", errors="ignore")
            except Exception as exc:
                raise HTTPException(
                    status_code=422,
                    detail=f"Could not parse PDF file: {exc}",
                )
        else:
            # .txt and other text files
            text = content.decode("utf-8", errors="ignore")

        text = text.strip()
        if not text:
            raise HTTPException(
                status_code=422,
                detail="No readable text found in the uploaded file.",
            )
        return {"text": text}
    except HTTPException:
        raise
    except Exception as exc:
        logger.error("File parsing failed: %s", exc)
        raise HTTPException(status_code=500, detail=str(exc))


@router.post("/generate-test-link", response_model=GenerateTestLinkResponse)
async def generate_test_link(
    request: GenerateTestLinkRequest,
    db: Session = Depends(get_db),
):
    """
    Create a test session and send the test link via email.
    Returns session_id and test_link.
    """
    # Get or create candidate
    candidate = get_or_create_candidate(
        db,
        CandidateCreate(
            email=request.email,
            name=request.name,
            candidate_id=request.candidate_id,
        ),
    )

    # Create the test session
    session = create_test_session(
        db,
        candidate=candidate,
        questions=request.questions,
        role=request.role,
        skills=request.skills,
        weights=request.weights,
    )

    # Send email
    email_sent = send_test_link_email(
        recipient_email=request.email,
        candidate_name=request.name,
        test_link=session.test_link,
        temp_password=session.temp_password,
        session_id=session.id,
    )

    if not email_sent:
        logger.warning("Email not sent for session %s", session.id)

    return GenerateTestLinkResponse(
        session_id=session.id,
        test_link=session.test_link,
        message=(
            "Test link sent via email."
            if email_sent
            else "Session created. Email service not configured — share the link manually."
        ),
    )


@router.post("/send-test-bulk")
async def send_test_bulk(
    request: BulkSendTestRequest,
    db: Session = Depends(get_db),
):
    """
    Send a test link to multiple candidates at once.
    """
    sent = 0
    results = []
    names = request.names or []

    for i, email in enumerate(request.emails):
        name = names[i] if i < len(names) else None
        try:
            candidate = get_or_create_candidate(
                db,
                CandidateCreate(email=email, name=name),
            )
            session = create_test_session(
                db,
                candidate=candidate,
                questions=request.questions,
                role=request.role,
                skills=request.skills,
                weights=request.weights,
            )

            # Link session to the job description if provided
            if request.jd_id:
                session.jd_id = request.jd_id
                db.commit()

            email_sent = send_test_link_email(
                recipient_email=email,
                candidate_name=name,
                test_link=session.test_link,
                temp_password=session.temp_password,
                session_id=session.id,
            )
            sent += 1
            results.append(
                {
                    "email": email,
                    "session_id": session.id,
                    "email_sent": email_sent,
                    "status": "success",
                }
            )
        except Exception as exc:
            logger.error("Bulk send failed for %s: %s", email, exc)
            results.append({"email": email, "status": "error", "detail": str(exc)})

    return {"sent": sent, "results": results}
